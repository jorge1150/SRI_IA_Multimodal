"""
tests/test_benchmark_dataset.py — Pruebas del parser de preguntas.docx.
Usa un .docx sintético (construido con python-docx) para no depender del
archivo real preguntas.docx del usuario. Ejecutar:
  python -m pytest tests/test_benchmark_dataset.py -v
"""

import os
import sys
import tempfile
import unittest

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)

from scripts.benchmark_dataset import parse_questions_docx


def _make_docx(path: str) -> None:
    from docx import Document
    doc = Document()
    doc.add_paragraph("").add_run("")
    p = doc.add_paragraph()
    p.add_run("Preguntas de Estudio - Normativa Tributaria SRI").bold = True
    doc.add_paragraph("2 Documentos PDF")
    doc.add_paragraph("")

    cat = doc.add_paragraph()
    cat.add_run("IVA (Impuesto al Valor Agregado)").bold = True
    doc.add_paragraph("")
    doc.add_paragraph("Documento A", style="Heading 2")
    doc.add_paragraph("Preguntas de comprensión y aplicación:", style="Heading 2")
    doc.add_paragraph("1. ¿Pregunta uno del documento A?")
    doc.add_paragraph("2. ¿Pregunta dos del documento A?")
    doc.add_paragraph("")

    doc.add_paragraph("Documento B (sin label, formato Pregunta N)")
    doc.add_paragraph("Pregunta 1: ¿Pregunta uno del documento B?")
    doc.add_paragraph("Pregunta 2: ¿Pregunta dos del documento B?")
    doc.save(path)


class TestParseQuestionsDocx(unittest.TestCase):

    def setUp(self):
        self._tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        self._tmp.close()
        _make_docx(self._tmp.name)

    def tearDown(self):
        if os.path.exists(self._tmp.name):
            os.unlink(self._tmp.name)

    def test_parses_all_questions(self):
        records = parse_questions_docx(self._tmp.name)
        self.assertEqual(len(records), 4)

    def test_category_and_source_doc_assigned(self):
        records = parse_questions_docx(self._tmp.name)
        for r in records:
            self.assertEqual(r["category"], "IVA (Impuesto al Valor Agregado)")
        self.assertEqual(records[0]["source_doc"], "Documento A")
        self.assertEqual(records[2]["source_doc"], "Documento B (sin label, formato Pregunta N)")

    def test_numbered_and_pregunta_format_both_parsed(self):
        records = parse_questions_docx(self._tmp.name)
        self.assertIn("¿Pregunta uno del documento A?", [r["question"] for r in records])
        self.assertIn("¿Pregunta uno del documento B?", [r["question"] for r in records])

    def test_title_and_doc_count_lines_ignored(self):
        records = parse_questions_docx(self._tmp.name)
        all_text = " ".join(r["question"] for r in records)
        self.assertNotIn("Preguntas de Estudio", all_text)
        self.assertNotIn("Documentos PDF", all_text)

    def test_label_line_never_becomes_source_doc(self):
        records = parse_questions_docx(self._tmp.name)
        for r in records:
            self.assertNotEqual(r["source_doc"], "Preguntas de comprensión y aplicación:")


if __name__ == "__main__":
    unittest.main()
